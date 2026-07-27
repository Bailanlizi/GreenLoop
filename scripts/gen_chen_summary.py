# -*- coding: utf-8 -*-
"""生成陈玉梅（成员C）个人实习总结，遵循模板结构与字数要求。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = "宋体"
TITLE_FONT = "黑体"


def set_run_font(run, size=12, bold=False, name=FONT):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_para(container, text, size=12, bold=False, name=FONT, align=None,
             space_after=6, line=1.5, first_indent=True):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if first_indent and not bold:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, name=name)
    return p


def add_heading(doc, text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, name=TITLE_FONT)
    return p


# ---------------- 正文内容 ----------------
title = "西南财经大学计算机与人工智能学院2026年夏季实训\n个人实习总结"

info_rows = [
    ("学　　号", "＿＿＿＿（请填写）"),
    ("班　　级", "＿＿＿＿（请填写）"),
    ("姓　　名", "陈玉梅"),
    ("实习项目", "GreenLoop 校园二手交易平台"),
    ("实习单位", "四川青软优加科技有限公司"),
    ("日　　期", "2026年7月14日 – 2026年7月21日"),
]

body = []

body.append(("一、学习内容和过程", None))
body.append((None,
    "本次夏季实训中，我们小组围绕“GreenLoop 校园二手交易平台”这一项目开展了为期约两周的设计与开发。"
    "作为小组组长兼成员C，我主要负责订单、配送、评价、信用分等交易主链路的后端开发，以及后期的联调测试与文档整理工作。"
    "下面从阅读文献、个人角色、技术实现和学习过程四个方面进行说明。"))
body.append(("1. 阅读的具体文献与资料", None))
body.append((None,
    "实训初期，为建立统一认知，我带领小组阅读并梳理了多份项目资料：产品需求文档（PRD）、团队《开发基准》、"
    "校园交易平台微服务架构设计文档、README 与 AI 部署自测说明；同时结合后端实现查阅了 Spring Boot 2.7、"
    "Spring Security、JWT、MyBatis、MySQL 8 与 Redis 的官方文档与社区资料，以及 Vue 3、Vite、Element Plus、"
    "Pinia 等前端技术资料。这些资料帮助我明确了校园二手交易的核心业务流程，也为后续模块设计提供了规范依据。"
    "此外，团队《开发基准》中的状态枚举、核心业务规则、测试验收与 Git 协作基准，以及微服务架构设计文档中关于"
    "前后端分离与三工程职责的说明，也是我反复查阅并严格遵循的依据；我还参考了状态机设计与事务一致性相关的技术文章，"
    "用于订单状态流转与信用分更新功能的实现。"))
body.append(("2. 我在本次项目实训中的角色", None))
body.append((None,
    "我在项目中承担双重职责。一方面作为后端负责人之一，独立负责订单、配送、评价、信用分模块的需求分析、"
    "接口设计、编码实现与单元/联调测试；另一方面作为组长，负责前期选题、任务分工、进度协调、需求分析汇总，"
    "并产出产品 PRD 文档，为全组开发提供统一目标与边界。这种角色让我既深入业务细节，也锻炼了项目协调能力。"))
body.append(("3. 采用何种技术实现", None))
body.append((None,
    "项目采用前后端分离的微服务结构，仓库包含三个工程：后端 API（campus-trade-api，端口 8080）、"
    "用户端（campus-trade-web，端口 5173）与管理端（campus-trade-admin，端口 8000）。"
    "后端以 Java 21 + Spring Boot 2.7.18 构建 REST 接口，使用 Spring Security + JWT 做无状态鉴权，"
    "MyBatis + MySQL 8 持久化，Redis 做缓存，PageHelper 分页，Spring Mail 发送邮件验证码。"
    "前端用户端与管理端均基于 Vue 3 + Vite + Vue Router + Pinia + Axios + Element Plus，"
    "管理端额外引入 ECharts 做数据统计。我负责的订单/评价/信用分模块即构建在这一技术基准之上，"
    "并严格遵循团队《开发基准》中的统一返回格式、认证权限边界与状态枚举约定。"
    "在接口层，后端统一使用 Result<T> 返回结构，通过 JwtAuthenticationFilter 解析令牌、在控制器中以"
    "@AuthenticationPrincipal 获取当前登录用户；订单状态变更与评价信用分更新均放在数据库事务中，保证数据一致性。"))
body.append(("4. 完成情况与学习过程", None))
body.append((None,
    "我的学习过程与开发推进基本同步，按时间线可分为以下阶段："))
body.append((None,
    "（1）7月14日，进行需求理解与任务拆分，梳理“商品浏览→提交订单→商品锁定→订单完成→用户评价→信用分变化”的"
    "完整交易闭环，完成订单、评价模块核心字段的初步设计，并识别需重点防范的异常场景（重复下单、购买自己商品、"
    "重复评价、越权改单）。"))
body.append((None,
    "（2）7月15日，明确订单模块核心接口（创建、我的购买、我的出售、详情、状态更新）与评价模块接口"
    "（提交、查询收到评价），整理订单创建校验规则，确认 orders、ratings 表与 user.credit_score 的关系，并统一字段命名。"))
body.append((None,
    "（3）7月16日，扩展订单模块以支持线下面交（MEETUP）与快递配送（SHIPPING）两种方式，将订单状态拆分为两条流转线："
    "线下面交 AWAITING_MEETUP→COMPLETED，快递配送 AWAITING_SHIPMENT→SHIPPED→COMPLETED。"))
body.append((None,
    "（4）7月17日，核对 OrderController / OrderServiceImpl / RatingController / RatingServiceImpl 代码，"
    "确认接口实现与重复评价唯一约束（uk_order_rater），并补充订单、评价测试用例。"))
body.append((None,
    "（5）7月18日，准备联调与测试数据（买家/卖家账号、商品、线下地点、收货地址、已完成订单），梳理完整交易流程测试顺序。"))
body.append((None,
    "（6）7月19日，对前期工作按当前实现复盘核对，补充异常测试点（重复下单、未完成订单评价、线下面交订单走快递发货、"
    "非订单用户操作等），明确验收重点。"))
body.append((None,
    "（7）7月20日，整理个人开发计划、开发日志、个人项目设计报告与实训报告，并按当前实际重写团队《开发基准》与"
    "《产品PRD》，将收藏、通知、求购、配送管理、AI 等已有功能纳入文档。"))
body.append((None,
    "（8）7月21日，配合完成《GreenLoop 校园二手交易平台》小组实习报告，负责订单、评价、信用分相关章节，"
    "并参与最终演示与验收说明。"))

body.append(("二、取得的成果", None))
body.append((None,
    "通过两周实训，我在订单、配送、评价、信用分这条交易主链路上取得了以下可验证成果："))
body.append((None,
    "第一，完成了订单模块的核心功能：创建订单时校验商品存在性、可购买状态、禁止购买自己商品，并保存成交价格；"
    "下单成功后商品状态由 AVAILABLE 变为 SOLD，取消订单时恢复为 AVAILABLE；提供我的购买、我的出售、订单详情与状态更新接口。"))
body.append((None,
    "第二，完成了评价与信用分模块：仅允许 COMPLETED 订单评价，同一用户对同一订单不可重复评价"
    "（数据库唯一约束保证），评分限制在 1–5，评价插入与信用分更新在同一事务内完成；信用分规则统一为"
    "5星+3、4星+1、3星0、2星-1、1星-3，最低不低于 0。"))
body.append((None,
    "第三，支持线下面交与快递配送两种模式，并设计了与之匹配的订单状态流转与权限控制"
    "（如线下面交订单不能走快递发货接口、快递订单仅 AWAITING_SHIPMENT 可发货）。"))
body.append((None,
    "第四，整理并跑通了完整交易闭环测试，覆盖正常流程与十余项异常场景（未登录下单、购买自己商品、商品不存在/"
    "已售/已下架、重复下单、越权改单、线下面交发货、未完成订单评价、评分越界、重复评价等），形成可复用的测试数据与用例。"))
body.append((None,
    "第五，产出个人与团队文档：个人开发计划、开发日志（7/14–7/20）、个人项目设计报告、实训报告，并重写团队"
    "《开发基准》与《产品PRD》，使文档准确反映当前实现而非早期方案；配合完成小组实习报告，保证资料一致、可提交。"))

body.append(("三、存在的问题", None))
body.append((None,
    "复盘中也暴露出一些问题。一是早期计划与当前实现存在差异：原计划强调卖家确认、取货码验证与“已预订（RESERVED）”"
    "状态，而当前项目以线下面交+快递配送为主，商品状态使用 AVAILABLE/SOLD/DELISTED，取货码与 RESERVED 状态尚未落地，"
    "需在文档中明确为后续优化方向。二是评价信用分规则在代码与文档间曾出现不一致（代码一度为高分+1、低分-2），"
    "后经统一修正，说明核心业务规则应尽早固化并同步到文档与页面。三是组内沟通不够及时、成员进度不够透明，"
    "后期常需通过代码与文档变化推断整体完成情况。四是后期文档较多、侧重点不同，若不及时区分用途容易出现内容重复或口径冲突。"))

body.append(("四、个人收获与体会", None))
body.append((None,
    "本次实训让我对“软件开发不仅是写接口”有了切身体会。过去我认为后端就是把功能跑通，但真实项目中必须同时考虑"
    "商品状态、订单状态、用户权限、重复下单、重复评价等异常，才能保证交易流程完整稳定。我也完整经历了从需求分析、"
    "功能设计、编码实现、联调测试到文档整理的软件工程全过程，对微服务前后端分离架构、Spring 安全鉴权、"
    "事务一致性与状态机设计有了更扎实的理解。"
    "例如，在核对代码时我发现若不加约束可能出现重复评价导致信用分被重复扣减的问题，后来通过数据库 uk_order_rater "
    "唯一约束与事务控制解决——这让我真正理解“异常场景”不是边界情况，而是业务正确性的底线。"))
body.append((None,
    "在团队层面，我体会到协作与沟通的重要性。作为组长，我认识到每日进度同步、任务记录与及时纠偏的价值；"
    "文档不只是交付物，更是团队对齐认知、降低沟通成本的工具。今后在团队项目中，我会更主动地推进规范化开发，"
    "更重视进度透明与接口变更通知。"))
body.append((None,
    "技术之外，这次实训也锻炼了我的文档编写、需求抽象与项目协调能力。后续我会继续加强业务建模、接口测试和团队沟通，"
    "在开发中更加主动、规范地推进工作，努力成长为既懂业务又懂协作的后端工程师。"))
body.append((None,
    "总体而言，这次夏季实训是一次从课堂走向工程实践的宝贵经历。它不仅让我把分散的技术知识点串联成完整的项目能力，"
    "也让我在协作与文档中学会如何与团队共同把一件事做实、做稳。我会把这次实训的收获带到今后的学习与工作中，持续进步。"))

# ---------------- 生成文档 ----------------
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

# 标题
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = Pt(12)
for i, line in enumerate(title.split("\n")):
    r = tp.add_run(line)
    set_run_font(r, size=18, bold=True, name=TITLE_FONT)
    if i == 0:
        tp.add_run("\n")

# 信息表
table = doc.add_table(rows=len(info_rows), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
for ri, (k, v) in enumerate(info_rows):
    c0 = table.cell(ri, 0)
    c0.text = ""
    add_para(c0, k, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=0)
    c1 = table.cell(ri, 1)
    c1.text = ""
    add_para(c1, v, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_after=0)
table.columns[0].width = Cm(3.5)
table.columns[1].width = Cm(11.0)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 正文
char_count = 0
for heading, text in body:
    if heading is not None:
        add_heading(doc, heading)
    if text is not None:
        add_para(doc, text)
        # 统计中文字符与标点（用于字数核对）
        for ch in text:
            if not ch.isspace():
                char_count += 1

print("正文字数（含标点，不含空白）:", char_count)
out = r"D:\Project\GreenLoop\个人文档\个人实习总结-陈玉梅.docx"
doc.save(out)
print("saved:", out)
