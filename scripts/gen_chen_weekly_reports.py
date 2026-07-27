# -*- coding: utf-8 -*-
"""生成陈玉梅（成员C）两份工作周报，遵循模板字段结构。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONT = "宋体"
TITLE_FONT = "黑体"


def set_run_font(run, size=12, bold=False, name=FONT):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_para(cell, text, size=12, bold=False, name=FONT, space_after=2):
    p = cell.paragraphs[0] if (len(cell.paragraphs) == 1 and not cell.paragraphs[0].text) else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, name=name)
    return p


def build_report(path, name, date_range, project, sections):
    doc = Document()
    # 默认正文字体
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    table = doc.add_table(rows=6, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 标题行
    title_cell = table.cell(0, 0).merge(table.cell(0, 5))
    title_cell.text = ""
    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("工作周报")
    set_run_font(tr, size=18, bold=True, name=TITLE_FONT)

    # 信息行：姓名 | 值 | 时间 | 值 | 团队项目名称 | 值
    info = [
        ("姓名", name, "时间", date_range, "团队项目名称", project),
    ]
    info_labels = ["姓名", "时间", "团队项目名称"]
    info_vals = [name, date_range, project]
    idx = 0
    for ci in range(6):
        c = table.cell(1, ci)
        c.text = ""
        if ci in (0, 2, 4):
            add_para(c, info_labels[ci // 2], bold=True, size=12)
        else:
            add_para(c, info_vals[(ci - 1) // 2], size=12)

    # 内容行
    sec_order = ["工作总结", "现存问题", "改进", "工作计划"]
    for ri, key in enumerate(sec_order):
        row = 2 + ri
        label_cell = table.cell(row, 0)
        label_cell.text = ""
        add_para(label_cell, key, bold=True, size=12)
        content_cell = table.cell(row, 1).merge(table.cell(row, 5))
        content_cell.text = ""
        for line in sections[key]:
            add_para(content_cell, line, size=12, space_after=3)

    # 列宽（参考模板：标签列窄，时间列较宽）
    widths = [Cm(2.2), Cm(2.6), Cm(1.8), Cm(4.2), Cm(2.6), Cm(2.6)]
    for ci, w in enumerate(widths):
        for row in table.rows:
            row.cells[ci].width = w

    doc.save(path)
    print("saved:", path)


base = r"D:\Project\GreenLoop\个人文档"
name = "陈玉梅"
project = "GreenLoop 校园二手交易平台"

# ---------- 第一周：2026-07-14 ~ 2026-07-18 ----------
week1 = {
    "工作总结": [
        "完成的内容：",
        "• 阅读产品 PRD 与开发计划，明确本人负责订单、评价、信用分及后期联调测试的任务范围。",
        "• 梳理完整交易闭环（商品浏览→提交订单→商品锁定→订单完成→用户评价→信用分变化），完成订单、评价模块核心字段初步设计（商品ID、买家/卖家ID、订单状态、成交价格、配送方式、评分、评价内容等）。",
        "• 设计订单模块核心接口：创建订单、查询我的购买/出售订单、订单详情、更新订单状态；评价模块接口：提交评价、查询用户收到的评价。",
        "• 整理订单创建校验规则（商品必须存在、状态可购买、禁止购买自己商品、下单后商品状态变更防重复购买）与信用分影响规则。",
        "• 与团队统一字段命名，确认 orders、ratings 表与 user.credit_score 的关系；确认登录用户ID获取方式，为订单接口鉴权做准备。",
        "• 扩展订单模块支持线下面交（MEETUP）与快递配送（SHIPPING）两种方式，明确不同方式初始状态（AWAITING_MEETUP / AWAITING_SHIPMENT）及必填字段（线下交易地点 / 收货地址）。",
        "• 检查 OrderController / OrderServiceImpl / RatingController / RatingServiceImpl 代码，确认已有订单、评价接口实现与重复评价唯一约束（uk_order_rater）。",
        "• 整理订单、评价、信用分相关测试数据（买家/卖家账号、可购买商品、线下交易地点、收货地址、已完成订单）与完整交易流程测试顺序。",
        "未完成的内容：",
        "• 信用分计算规则尚未与文档完全统一（代码当前为高分+1/低分-2，文档建议5星+3至1星-3）。",
        "• 评分范围（1~5）的参数校验尚未补全。",
        "• 用户端“我的购买/我的出售/评价弹窗”与管理端订单、配送页面尚未完成前后端联调。",
    ],
    "现存问题": [
        "• 早期计划的部分状态命名与当前代码不一致（如商品状态由 ON_SALE/RESERVED/SOLD 调整为 AVAILABLE/SOLD/DELISTED；订单状态需兼容待发货、已发货），需在文档中明确以当前实现为准。",
        "• 配送流程加入后，订单状态比原计划复杂，需拆分为线下面交与快递配送两条流转线。",
        "• 部分文档与 SQL 注释在命令行中显示乱码（IDE 中正常），暂以 IDE 显示效果为准。",
    ],
    "改进": [
        "• 订单流程按两条线拆分：线下面交 AWAITING_MEETUP→COMPLETED；快递配送 AWAITING_SHIPMENT→SHIPPED→COMPLETED，避免状态设计过于单一。",
        "• 商品状态以当前代码 AVAILABLE/SOLD/DELISTED 为准，并在文档中说明 SOLD 当前同时表示已售出和已被订单占用。",
        "• 后续统一 RatingServiceImpl 信用分计算规则，并补充评分范围校验，防止出现 0 分或超过 5 分的异常评分。",
        "• 文档统一使用 UTF-8 保存，命令行乱码不作为功能问题处理。",
    ],
    "工作计划": [
        "• 继续检查 OrderServiceImpl 中订单状态修改的合法性，明确哪些状态可取消、哪些不可取消。",
        "• 统一 RatingServiceImpl 信用分计算规则为 5星+3 / 4星+1 / 3星0 / 2星-1 / 1星-3，并补充评分 1~5 校验。",
        "• 联调用户端“我的购买/我的出售/评价弹窗”页面与管理端订单、配送管理功能。",
        "• 准备买家、卖家、商品、订单、评价演示数据，跑通完整交易闭环。",
        "• 整理订单状态流转图与信用分规则说明，为答辩做准备。",
    ],
}

# ---------- 第二周：2026-07-19 ~ 2026-07-21 ----------
week2 = {
    "工作总结": [
        "完成的内容：",
        "• 对 7/14–7/18 完成的订单、配送、评价、信用分工作按当前项目实际情况进行复盘核对：确认商品 AVAILABLE 可下单、下单后变 SOLD、取消订单恢复 AVAILABLE；确认订单状态枚举（AWAITING_MEETUP / AWAITING_SHIPMENT / SHIPPED / COMPLETED / CANCELLED）。",
        "• 补充异常测试点：重复下单、购买自己的商品、未完成订单评价、重复评价、线下面交订单走快递发货、非订单相关用户操作订单等，并明确预期拒绝结果。",
        "• 检查用户端订单相关页面与接口封装（我的购买/我的出售/订单状态更新），确认与管理端订单、配送页面（发货、统计、导出）保持一致。",
        "• 整理个人负责模块验收重点：交易流程能跑通、异常规则能说清楚、信用分规则能解释清楚。",
        "• 整理个人开发计划、开发日志（7/14–7/20 时间线）、个人项目设计报告（订单/配送/评价/信用分模块设计）与实训报告。",
        "• 根据当前项目实际情况重写开发基准.md 与产品PRD.md，纳入收藏、通知、求购、配送管理、AI 等已有功能，去除与当前工程不符的旧内容。",
        "• 新建“个人文档”文件夹归档个人材料，更新 .gitignore；整理当前项目与早期 PRD 差异，明确已落地与后续优化功能。",
        "• 配合完成《GreenLoop 校园二手交易平台》小组实习报告，提供订单、评价、信用分相关章节（详细设计、测试结果）的内容。",
        "未完成的内容：",
        "• 取货码验证流程与商品 RESERVED 状态当前项目未完整落地，列为后续优化方向，未在报告中写成已实现。",
        "• 用户端“我的购买/我的出售”评价弹窗联调与部分管理端配送联调仍需在最终演示前再次确认。",
    ],
    "现存问题": [
        "• 早期计划与当前实现存在差异（原计划强调卖家确认、取货码、“已预订”状态；当前以线下面交+快递配送为主，商品状态为 AVAILABLE/SOLD/DELISTED），需以当前实际为准进行验收说明。",
        "• 后期文档较多、侧重点不同，易出现内容重复或口径不一致（设计报告偏模块设计、实训报告偏实训经历、开发基准偏规范、PRD 偏需求）。",
    ],
    "改进": [
        "• 以后期项目实际情况为准进行验收和说明，未落地的取货码/RESERVED 状态在文档中明确为“后续可优化方向”，不强行写成已实现。",
        "• 按文档用途区分内容：个人开发计划（任务安排）、开发日志（每日记录）、个人项目设计报告（模块设计）、实训报告（实训经历）、开发基准（统一规范）、PRD（产品需求），避免重复与口径冲突。",
        "• 重点保证已有订单、配送、评价、信用分流程稳定，最终演示聚焦交易流程完整性、异常规则清晰度、信用分变化合理性。",
    ],
    "工作计划": [
        "• 继续检查 OrderServiceImpl 订单状态修改合法性，统一 RatingServiceImpl 信用分计算规则并补充评分校验。",
        "• 联调用户端“我的购买/我的出售/评价弹窗”与管理端订单、配送管理功能，准备演示数据。",
        "• 整理答辩中订单状态流转图与信用分规则说明。",
        "• 配合小组完成最终演示、验收说明与材料提交。",
    ],
}

build_report(base + r"\工作周报-陈玉梅-第一周.docx", name, "2026年7月14日--2026年7月18日", project, week1)
build_report(base + r"\工作周报-陈玉梅-第二周.docx", name, "2026年7月19日--2026年7月21日", project, week2)
print("ALL DONE")
