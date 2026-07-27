# -*- coding: utf-8 -*-
"""生成《GreenLoop 校园二手交易平台》小组实习报告 Word 文档。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Project\GreenLoop\GreenLoop校园二手交易平台_小组实习报告.docx"

FONT_CN = "宋体"
FONT_CN_H = "黑体"

doc = Document()

# ------- 基础样式：正文宋体小四，行距 --------
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

def set_cn_font(run, cn=FONT_CN, size=None, bold=False, color=None):
    run.font.name = "Times New Roman"
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), cn)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def para(text="", size=12, cn=FONT_CN, bold=False, align=None, after=6, before=0, indent_first=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = 1.5
    if indent_first and text:
        pf.first_line_indent = Pt(size * 2)
    if text:
        run = p.add_run(text)
        set_cn_font(run, cn, size, bold)
    return p

def h1(num_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(num_text)
    set_cn_font(run, FONT_CN_H, 16, True)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_cn_font(run, FONT_CN_H, 14, True)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_cn_font(run, FONT_CN_H, 12.5, True)
    return p

def bullets(items, size=12):
    for it in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run("● " + it)
        set_cn_font(run, FONT_CN, size)

def add_table(headers, rows, widths=None, size=10.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        pr = hdr[i].paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pr.add_run(htext)
        set_cn_font(run, FONT_CN_H, size, True)
        # 表头底纹
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E8DC")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            pr = cells[i].paragraphs[0]
            pr.paragraph_format.line_spacing = 1.15
            run = pr.add_run(str(val))
            set_cn_font(run, FONT_CN, size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================================================
# 封面
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("小 组 实 习 报 告"); set_cn_font(r, FONT_CN_H, 30, True)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("GreenLoop 校园二手交易平台"); set_cn_font(r, FONT_CN_H, 20, True, RGBColor(0x1F, 0x6B, 0x3B))

for _ in range(3):
    doc.add_paragraph()

info = [
    ("设 计 题 目", "GreenLoop 校园二手交易平台"),
    ("班        级", "________________"),
    ("小 组 成 员", "田敬萱、余夏蓉、陈玉梅"),
    ("学        号", "________________"),
    ("完 成 日 期", "2026 年 7 月 21 日"),
]
t = doc.add_table(rows=len(info), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info):
    c0 = t.rows[i].cells[0]; c1 = t.rows[i].cells[1]
    c0.width = Cm(4.5); c1.width = Cm(8.5)
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run(k); set_cn_font(run0, FONT_CN_H, 14, True)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(v); set_cn_font(run1, FONT_CN, 14)
# 去除封面信息表边框
tbl = t._tbl
for cell in t._cells:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tcPr.append(borders)

doc.add_page_break()

# ============================================================
# 摘要小节（可选说明）
# ============================================================
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("GreenLoop 校园二手交易平台 · 小组实习报告"); set_cn_font(r, FONT_CN_H, 16, True)
para("本报告为软件项目实习小组共同完成的成果材料，围绕 GreenLoop 校园二手交易平台的"
     "需求分析、概要设计、详细设计、测试结果等内容展开，完整记录了小组从选题立项、"
     "任务分工、系统设计到编码实现、联调测试的全过程。平台采用 Vue 3 + Spring Boot "
     "前后端分离架构，覆盖用户、商品、消息、订单、配送、评价、信用分及后台管理等核心业务，"
     "形成了一条可演示、可追踪的完整校园二手交易闭环。", after=10)

# ============================================================
# 一、项目介绍
# ============================================================
h1("一、项目介绍")

h2("1. 项目目的与要求")
para("随着高校学生规模的扩大和消费更新频率的加快，教材、电子产品、生活用品、文具、代步工具等"
     "闲置物品在校园内大量产生。传统的二手流转方式主要依赖微信群、QQ 群和校园信息墙，普遍存在"
     "信息分散、历史消息容易被覆盖、商品描述与分类不统一、买卖双方缺乏基础信用信息、交易过程"
     "无标准订单记录等问题，导致交易效率低、体验差、纠纷难追溯。")
para("本项目的目的，是设计并实现一个面向高校学生的校园二手交易平台 GreenLoop，为学生提供结构化、"
     "易检索、可追踪的二手交易服务，围绕“闲置物品流转”这一真实校园场景，打通从商品发布、浏览搜索、"
     "在线沟通、下单交易、配送处理到评价反馈、信用积累的完整业务闭环，并配套后台运营管理能力。")
para("项目的主要要求如下：")
bullets([
    "采用前后端分离架构，前端负责页面与交互，后端负责业务逻辑、数据持久化与安全认证；",
    "实现用户注册登录与基于 JWT 的鉴权，保障受保护接口的访问安全；",
    "实现商品发布、编辑、下架、搜索筛选与详情展示，并支持图片上传；",
    "实现站内私信、收藏、求购、通知等社交与辅助功能；",
    "实现订单创建、状态流转、线下面交与快递配送两种配送方式；",
    "实现订单评价与用户信用分体系，形成交易反馈机制；",
    "提供管理后台，支持用户、商品、订单、配送、交易地点管理与数据统计；",
    "保证至少一条完整交易流程可稳定演示，异常场景处理正确，接口返回格式统一。",
])

h2("2. 项目任务计划")
para("小组开发周期为 2026 年 7 月 14 日至 2026 年 7 月 20 日，按“需求分析 → 系统设计 → 编码实现 → "
     "联调测试 → 文档整理”的阶段推进，各阶段主要任务与交付物安排如下：")
add_table(
    ["阶段", "时间", "主要任务", "交付物"],
    [
        ["需求与设计", "07-14 至 07-15", "选题立项、需求分析、用户角色与用例梳理，确定技术栈与数据库字段、接口清单", "产品 PRD、开发基准、接口与数据表约定"],
        ["核心开发", "07-16 至 07-17", "用户/商品/消息、订单/评价/信用等后端接口实现，前端页面搭建与联调准备", "各模块后端接口、前端页面框架"],
        ["联调测试", "07-18 至 07-19", "前后端与管理端联调，交易闭环与异常场景测试，问题修复", "可跑通的交易流程、测试记录、Bug 清单"],
        ["整理归档", "07-20 至 07-21", "演示数据准备、文档整理、报告撰写与成果归档", "演示数据、开发日志、实习报告"],
    ],
    widths=[3.0, 3.2, 6.5, 4.5],
)

h2("3. 项目成员介绍与分工")
para("本小组由三名成员组成，按照“前端 + 两条后端主线”的方式进行分工，既保证模块职责清晰，又便于"
     "接口对接与联调协作。具体分工如下：")
add_table(
    ["姓名", "角色定位", "主要职责"],
    [
        ["田敬萱", "前端与页面负责人",
         "负责用户端与管理端所有页面开发，包括登录注册、首页商品流、商品详情、发布/编辑商品、订单确认、"
         "个人中心、消息中心、聊天窗口，以及管理端仪表盘与各管理页面；负责 API 封装、路由拦截、登录态管理与页面交互体验。"],
        ["余夏蓉", "用户、商品与消息后端负责人",
         "负责用户注册登录与 JWT 鉴权、个人资料、邮箱验证码；商品发布/编辑/下架、分类、搜索筛选、图片上传、"
         "收藏、求购与商品推荐；站内私信与系统通知等后端接口的设计与实现。"],
        ["陈玉梅", "组长 / 订单、评价与信用后端负责人",
         "担任组长，负责前期需求分析与产品 PRD 编写、任务分工与进度协调；开发阶段负责订单创建与状态流转、"
         "线下面交与快递配送、订单评价与信用分规则，以及系统测试、演示数据准备与答辩材料整理。"],
    ],
    widths=[2.2, 4.0, 10.5],
)

# ============================================================
# 二、项目内容
# ============================================================
h1("二、项目内容")

# ---------- 2.1 需求分析 ----------
h2("1. 需求分析")

h3("1.1 目标用户与角色")
para("平台面向校园二手交易场景，主要涉及三类角色，各角色的核心诉求如下：")
add_table(
    ["角色", "核心诉求"],
    [
        ["买家", "搜索低价二手商品、查看商品详情、与卖家沟通并下单，通过信用分判断交易对象可靠度"],
        ["卖家", "发布闲置商品、处理订单、完成交易并积累信用"],
        ["管理员", "管理平台用户、商品、订单、配送与交易地点，查看运营统计数据"],
    ],
    widths=[3.0, 13.0],
)

h3("1.2 功能性需求")
para("依据用户角色与业务闭环，梳理出以下核心功能性需求：")
add_table(
    ["功能模块", "功能需求描述"],
    [
        ["用户与认证", "用户名唯一、密码加密存储；注册、登录、JWT 鉴权；个人资料维护；邮箱验证码"],
        ["商品管理", "商品发布/编辑/下架；固定分类；成色、价格、封面图等字段；我的商品管理"],
        ["搜索与推荐", "关键词搜索、分类与价格区间筛选、排序、相似商品推荐、AI 语义搜索与混合推荐"],
        ["社交与辅助", "站内私信、商品收藏、求购发布、系统通知"],
        ["订单交易", "创建订单、查询我的购买/出售、订单状态流转、下单锁定商品防重复购买"],
        ["配送管理", "线下面交（选择交易地点）与快递配送（选择收货地址、填写物流并发货）"],
        ["评价与信用", "已完成订单评价、评分 1~5、防重复评价、评价后更新被评价用户信用分"],
        ["后台管理", "用户/商品/订单/配送/交易地点管理，配送导出与平台数据统计"],
    ],
    widths=[3.4, 12.6],
)

h3("1.3 非功能性需求")
bullets([
    "安全性：基于 Spring Security + JWT 的无状态认证，密码使用 BCrypt 加密，接口按角色控制权限；",
    "一致性：商品状态与订单状态必须联动一致，评价写入与信用分更新在同一事务中完成，防止重复变更；",
    "性能：使用 Redis 缓存分类等热点数据，商品列表分页查询，减少数据库压力；",
    "可用性：接口统一返回 Result 结构，前端对空状态、加载状态与错误提示进行统一处理；",
    "可维护性：分层架构、目录规范统一，接口路径、状态枚举、数据库字段随代码同步维护。",
])

h3("1.4 核心业务闭环")
para("平台围绕以下核心交易闭环组织全部功能，这也是需求分析中最重要的主线：")
para("用户注册登录 → 卖家发布商品 → 买家浏览搜索 → 私信沟通 → 提交订单（锁定商品）→ "
     "线下面交或快递配送 → 订单完成 → 用户评价 → 信用分更新 → 管理端查看与管理。",
     cn=FONT_CN, bold=True, indent_first=False)

# ---------- 2.2 概要设计 ----------
h2("2. 概要设计")

h3("2.1 设计思想")
para("概要设计以“业务闭环驱动、职责单一、前后端解耦、状态一致”为核心思想：")
bullets([
    "以交易闭环为主线组织系统：所有模块围绕“从商品发布到订单完成再到评价反馈”的主流程展开，保证核心链路完整、可演示；",
    "前后端分离、契约优先：前端与后端通过统一的 REST 接口和 Result 返回结构协作，接口路径、字段命名先约定后实现；",
    "模块高内聚、低耦合：按业务领域划分用户、商品、消息、订单、配送、评价、后台管理等模块，各模块内聚，模块间通过明确的数据关系与接口交互；",
    "状态机思想管理生命周期：商品与订单均采用有限状态枚举，状态迁移只能沿合法路径进行，避免出现不一致的中间态；",
    "安全与权限内建：默认接口需登录，仅少量浏览类接口放行，用户只能操作属于自己的资源，管理端接口限管理员访问；",
    "数据一致性优先：涉及跨表更新（如下单改商品状态、评价改信用分）的操作放入事务，保证要么全部成功、要么全部回滚。",
])

h3("2.2 实现方法（技术选型与架构）")
para("系统采用主流的前后端分离技术栈，仓库以 Monorepo 方式统一管理后端 API、用户端与管理端三个工程。"
     "各工程的技术选型如下：")
add_table(
    ["工程", "路径", "职责", "技术栈", "端口"],
    [
        ["后端 API", "campus-trade-api", "提供全部 REST 接口、业务逻辑与数据持久化",
         "Java 21、Spring Boot 2.7.18、Spring Security、JWT、MyBatis、MySQL 8、Redis、PageHelper、Spring Mail", "8080"],
        ["用户端", "campus-trade-web", "面向普通用户的浏览、发布、下单、私信等功能",
         "Vue 3、Vite、Vue Router、Pinia、Axios、Element Plus", "5173"],
        ["管理端", "campus-trade-admin", "面向管理员的用户、商品、订单、配送、统计管理",
         "Vue 3、Vite、Vue Router、Pinia、Axios、Element Plus、ECharts", "8000"],
    ],
    widths=[2.4, 3.4, 4.2, 4.6, 1.4],
)
para("后端遵循 controller → service → mapper 分层架构：控制器负责接收请求与鉴权，服务层承载业务逻辑，"
     "Mapper 接口配合 XML 完成数据库操作；使用 Spring Security + JwtAuthenticationFilter 解析令牌并注入当前登录用户；"
     "接口统一返回 Result<T>（code / msg / data）结构，前端以 code == 200 判断业务成功。")

h3("2.3 系统主要模块及各模块间关系")
para("系统在逻辑上划分为用户端、后端服务与管理端三层，后端服务内部按业务领域进一步划分为若干模块。"
     "主要模块及其职责如下：")
add_table(
    ["模块", "主要职责"],
    [
        ["用户模块", "注册登录、JWT 鉴权、个人资料、信用分存储（credit_score）"],
        ["商品模块", "商品 CRUD、分类、搜索筛选、图片上传、收藏、求购、推荐"],
        ["消息模块", "站内私信、会话列表、系统通知"],
        ["订单模块", "订单创建、状态流转、我的购买/出售、下单锁定商品"],
        ["配送模块", "线下面交地点、收货地址、快递发货、配送统计与导出"],
        ["评价与信用模块", "订单评价、评分校验、防重复评价、信用分更新"],
        ["后台管理模块", "用户/商品/订单/配送/地点管理与平台数据统计"],
    ],
    widths=[3.4, 12.6],
)
para("各模块并非孤立存在，而是围绕订单这一核心业务相互关联。主要关系如下：")
add_table(
    ["关联关系", "说明"],
    [
        ["用户 ↔ 商品", "用户作为卖家发布商品（user 1 — n product）"],
        ["用户 ↔ 订单", "订单包含买家与卖家两个用户，均来自用户模块"],
        ["商品 ↔ 订单", "创建订单需读取商品信息并将商品状态改为 SOLD（product 1 — n orders）"],
        ["订单 ↔ 配送", "订单按配送方式关联交易地点或收货地址，快递订单可发货"],
        ["订单 ↔ 评价", "订单完成后产生评价（orders 1 — n ratings）"],
        ["评价 ↔ 用户", "评价结果回写被评价用户的信用分"],
        ["消息 ↔ 商品/用户", "买卖双方围绕商品进行私信沟通"],
        ["各模块 ↔ 后台管理", "管理端对用户、商品、订单、配送、地点进行统一管理与统计"],
    ],
    widths=[3.6, 12.4],
)
para("由此可见，订单模块是连接用户、商品、配送、评价与后台管理的枢纽，商品状态与订单状态的联动是"
     "保证系统数据一致性的关键。", indent_first=True)

h3("2.4 数据库概要设计")
para("系统采用 MySQL 作为主数据库，初始化脚本为 campus-trade-api/src/main/resources/campus_trade.sql。"
     "核心数据表及其说明如下：")
add_table(
    ["数据表", "说明"],
    [
        ["user", "用户表，含用户名、密码、昵称、角色、状态、信用分等"],
        ["category", "商品分类表"],
        ["product", "商品表，含卖家、分类、价格、成色、状态、支持配送方式"],
        ["product_images", "商品图片表"],
        ["orders", "订单表，含商品、买卖双方、订单状态、成交价、配送信息"],
        ["ratings", "评价表，含订单、评价人、被评价人、评分、内容"],
        ["messages", "私信表"],
        ["favorites", "收藏表"],
        ["notifications", "通知表"],
        ["meetup_location", "线下交易地点表"],
        ["user_addresses", "用户收货地址表"],
        ["product_demands", "求购需求表"],
        ["product_embeddings / product_risks", "商品向量与风险记录表（AI 相关）"],
    ],
    widths=[5.2, 10.8],
)
para("表间关系概括为：user 1—n product、user 1—n orders、product 1—n orders、orders 1—n ratings、"
     "user 1—n ratings、user 1—n user_addresses、meetup_location 1—n orders。", indent_first=True)

# ---------- 2.3 详细设计 ----------
h2("3. 详细设计")
para("本节按小组分工，分别说明三位成员负责模块的详细设计，包括关键字段、接口、状态流转与核心规则。")

h3("3.1 前端与页面详细设计（田敬萱）")
para("前端分为用户端（campus-trade-web）与管理端（campus-trade-admin），均采用 Vue 3 Composition API + "
     "Element Plus 组件化开发，接口统一封装于 src/api，路由集中管理并对需登录页面做拦截，登录态与用户信息由 Pinia 维护。")
para("主要页面清单如下：")
add_table(
    ["端", "页面 / 组件", "说明"],
    [
        ["用户端", "Login.vue / Register.vue", "登录、注册（邮箱验证码）"],
        ["用户端", "Home.vue", "首页商品流、分类与价格筛选、排序"],
        ["用户端", "ProductDetail.vue", "商品详情、卖家信息与信用分、下单入口"],
        ["用户端", "PublishProduct.vue / EditProduct.vue", "发布/编辑商品，多图上传、AI 润色"],
        ["用户端", "OrderConfirmation.vue", "订单确认，选择配送方式与地址/地点"],
        ["用户端", "Dashboard.vue", "个人中心：我的商品/购买/出售/收藏/通知/地址/求购、评价弹窗"],
        ["用户端", "messages/index.vue、ChatWindow.vue", "消息中心与聊天窗口"],
        ["管理端", "dashboard/index.vue", "数据仪表盘（ECharts 可视化）"],
        ["管理端", "user / product / order / delivery / location", "用户、商品、订单、配送、交易地点管理"],
    ],
    widths=[2.0, 5.4, 8.6],
)
para("前端交互要点：表单提交前做基础校验；异步请求提供加载状态；空列表显示空状态；后端业务错误通过 msg 明确提示；"
     "订单、商品、评价等按钮根据当前状态动态禁用或隐藏；接口返回未登录或令牌失效时清理登录态并跳转登录页。")

h3("3.2 用户、商品与消息模块详细设计（余夏蓉）")
para("该部分承载平台的“人”与“物”以及沟通能力，是交易闭环的前半段。")
para("① 用户与认证：注册接口 POST /users/register，登录接口 POST /users/authenticate，登录成功返回 JWT；"
     "GET /users/{id}/profile 查看公开资料，PUT /api/me/profile 更新资料；密码加密存储，用户名唯一。")
para("② 商品与搜索：")
add_table(
    ["功能", "接口"],
    [
        ["商品列表 / 搜索 / 筛选", "GET /products"],
        ["商品详情", "GET /products/{id}"],
        ["发布 / 编辑 / 改状态", "POST /products、PUT /products/{id}、PUT /products/{id}/status"],
        ["我的商品", "GET /products/my"],
        ["相似推荐 / 风险检测", "GET /products/{id}/recommendations、POST /products/risk-check"],
        ["分类列表", "GET /categories"],
        ["图片上传", "POST /files/upload（单文件≤10MB，访问路径 /uploads/{filename}）"],
    ],
    widths=[5.0, 11.0],
)
para("商品状态采用 AVAILABLE（可购买）、SOLD（已售出/被订单占用）、DELISTED（已下架）三态；发布商品必须登录，"
     "标题非空、价格≥0、包含封面图，用户只能编辑自己的商品。")
para("③ 收藏、求购、私信与通知：收藏（POST/DELETE /api/products/{id}/favorite、GET /api/me/favorites）、"
     "求购（POST /demands、GET /demands/my、DELETE /demands/{id}）、私信（POST /api/messages、"
     "GET /api/messages/history/{otherUserId}、GET /api/messages/conversations）、"
     "通知（GET /api/notifications、/unread-count、/mark-all-as-read）。私信发送人与接收人不能相同、内容不能为空。")

h3("3.3 订单、评价与信用模块详细设计（陈玉梅）")
para("该部分位于交易闭环后半段，是平台从“商品展示”进入“真实交易”的关键，直接决定商品能否被正确购买、"
     "订单状态是否一致、交易完成后能否形成信用反馈。")
para("① 订单核心字段：id、productId、buyerId、sellerId、orderStatus、totalPrice、deliveryMethod、"
     "meetupLocationId、shippingAddressId、shippingProvider、trackingNumber、createTime。")
para("② 订单与配送接口：")
add_table(
    ["功能", "接口", "权限"],
    [
        ["创建订单", "POST /orders", "登录用户"],
        ["订单详情", "GET /orders/{orderId}", "订单相关用户"],
        ["我的购买 / 我的出售", "GET /orders/my-purchases、/my-sales", "登录用户"],
        ["更新订单状态", "PUT /orders/{orderId}/status", "订单相关用户"],
        ["提交评价", "POST /api/orders/{orderId}/ratings", "订单相关用户"],
        ["用户评价列表", "GET /api/users/{userId}/ratings", "公开/登录"],
        ["后台发货", "PUT /admin/orders/{id}/ship", "管理员"],
        ["配送统计 / 导出", "GET /admin/delivery/stats、/export", "管理员"],
    ],
    widths=[4.0, 8.5, 3.5],
)
para("③ 订单状态与状态流转：订单状态为 AWAITING_MEETUP、AWAITING_SHIPMENT、SHIPPED、COMPLETED、CANCELLED；"
     "配送方式 MEETUP（必填交易地点）与 SHIPPING（必填收货地址）。商品状态与订单状态联动如下：")
add_table(
    ["场景", "商品状态", "订单状态"],
    [
        ["商品发布后", "AVAILABLE", "无订单"],
        ["线下面交下单", "SOLD", "AWAITING_MEETUP"],
        ["快递配送下单", "SOLD", "AWAITING_SHIPMENT"],
        ["买家/管理员取消订单", "AVAILABLE", "CANCELLED"],
        ["快递订单发货", "SOLD", "SHIPPED"],
        ["订单完成", "SOLD", "COMPLETED"],
    ],
    widths=[5.0, 5.5, 5.5],
)
para("④ 订单创建校验：商品必须存在且为 AVAILABLE；买家不能购买自己的商品；配送方式必须被商品支持；"
     "线下面交必须选交易地点、快递必须选收货地址；下单成功后保存成交价格并将商品置为 SOLD 以防重复下单。")
para("⑤ 评价与信用分：只有 COMPLETED 订单可评价，评分限制 1~5，同一用户对同一订单仅能评价一次"
     "（数据库 uk_order_rater 唯一约束保证）。信用分规则统一为：")
add_table(
    ["星级", "信用分变化", "星级", "信用分变化"],
    [
        ["5 星", "+3", "2 星", "-1"],
        ["4 星", "+1", "1 星", "-3"],
        ["3 星", "0", "说明", "信用分最低不低于 0"],
    ],
    widths=[2.6, 3.4, 2.6, 3.4],
)
para("评价写入与信用分更新必须在同一事务中完成，重复提交不会导致信用分重复变化，从而保证信用体系的可靠性。")

# ---------- 2.4 测试结果 ----------
h2("4. 测试结果")

h3("4.1 测试环境")
add_table(
    ["项目", "配置"],
    [
        ["操作系统", "Windows"],
        ["后端运行环境", "JDK 21、Maven、Spring Boot 2.7.18"],
        ["数据库 / 缓存", "MySQL 8.x、Redis 6.x+"],
        ["前端运行环境", "Node.js 16+、Vite"],
        ["访问地址", "后端 8080、用户端 5173、管理端 8000"],
        ["管理员账号", "admin / admin123（系统启动自动创建）"],
    ],
    widths=[4.0, 12.0],
)

h3("4.2 功能测试（正常流程）")
add_table(
    ["用例", "预期结果", "结果"],
    [
        ["用户注册并登录", "注册成功，登录返回 JWT，可访问受保护接口", "通过"],
        ["发布商品并上传图片", "商品创建成功，图片可正常访问", "通过"],
        ["搜索与筛选商品", "关键词、分类、价格筛选与排序结果正确", "通过"],
        ["买家对可售商品下单", "订单创建成功，商品状态变为 SOLD", "通过"],
        ["线下面交下单", "订单状态为 AWAITING_MEETUP", "通过"],
        ["快递配送下单并发货", "下单为 AWAITING_SHIPMENT，发货后为 SHIPPED", "通过"],
        ["订单完成", "订单状态变为 COMPLETED", "通过"],
        ["已完成订单评价", "评价保存成功，信用分按星级正确变化", "通过"],
        ["收藏 / 私信 / 通知", "收藏与取消、私信收发、通知未读数正常", "通过"],
        ["管理端查看订单与统计", "订单列表、配送信息与统计数据正确展示", "通过"],
    ],
    widths=[5.5, 8.5, 2.0],
)

h3("4.3 异常测试")
add_table(
    ["异常场景", "预期结果", "结果"],
    [
        ["未登录创建订单", "拒绝", "通过"],
        ["购买自己发布的商品", "拒绝", "通过"],
        ["对已售出/下架商品下单", "拒绝", "通过"],
        ["同一商品重复下单", "拒绝", "通过"],
        ["线下面交未选交易地点", "拒绝", "通过"],
        ["快递订单未选收货地址", "拒绝", "通过"],
        ["线下面交订单执行快递发货", "拒绝", "通过"],
        ["未完成订单提交评价", "拒绝", "通过"],
        ["同一订单重复评价", "拒绝，信用分不重复变化", "通过"],
        ["评分小于 1 或大于 5", "拒绝", "通过"],
        ["非订单相关用户操作订单", "拒绝", "通过"],
    ],
    widths=[6.0, 8.0, 2.0],
)

h3("4.4 交易闭环联调测试")
para("最终演示前完成了完整交易闭环联调，测试顺序与结果为：卖家注册登录 → 发布商品 → 买家注册登录 → "
     "搜索并查看详情 → 私信联系卖家 → 选择配送方式下单 → 商品变为不可重复购买 → 卖家/管理员处理订单 → "
     "订单完成 → 买家提交评价 → 卖家信用分正确变化 → 管理端查看订单与统计数据。全流程跑通，"
     "状态流转一致，异常场景均被正确拦截。")

h3("4.5 测试结论")
para("测试用例覆盖了交易流程的主要正常路径与关键异常路径。结果表明：订单模块能够正确连接商品、用户、"
     "配送与评价模块；评价模块能够在订单完成后形成信用反馈；管理端能够对订单与配送进行辅助管理；"
     "系统整体满足验收标准，核心交易闭环稳定可演示。后续可进一步引入自动化测试以降低手动联调成本。")

# ============================================================
# 三、项目总结
# ============================================================
h1("三、项目总结")
para("经过一周的集中开发，小组共同完成了 GreenLoop 校园二手交易平台的设计与实现。项目基于 Vue 3 + "
     "Spring Boot 前后端分离架构，覆盖用户、商品、消息、订单、配送、评价、信用分以及后台管理等核心业务，"
     "实现了一条从商品发布、浏览搜索、在线沟通、下单交易、配送处理到评价反馈、信用积累的完整校园二手交易闭环，"
     "并集成了 AI 文案润色、价格建议、语义搜索与混合推荐等增强能力。")
para("在技术上，小组实践了分层架构、Spring Security + JWT 鉴权、MyBatis 持久化、Redis 缓存、"
     "统一 Result 返回结构，以及 Vue 3 组件化开发与 Pinia 状态管理；在业务上，深入理解了商品状态与订单状态的"
     "联动、事务一致性、权限边界与异常防护等关键问题。特别是订单、评价与信用这条主链路，让我们认识到后端开发"
     "并非简单的接口拼装，而需要统筹商品状态、用户身份、配送方式、评价资格与信用变化等多重约束。")
para("在协作方面，小组按照“前端 + 两条后端主线”的方式分工，明确了各自职责与接口契约，并通过统一的开发基准文档"
     "约定接口路径、状态枚举与数据库字段，减少了联调冲突。同时也认识到，接口字段与状态枚举一旦变化必须及时同步"
     "前端与管理端，跨模块字段（如订单状态）尤其需要提前在文档中约定清楚；此外，日常进度同步与任务记录的及时性"
     "仍有提升空间。")
para("受实习周期限制，项目仍有可优化之处，例如：新增 RESERVED 状态以区分“已占用”与“已售出”；引入正式订单状态机"
     "限制非法状态跳转；将信用分规则进一步封装；补充更多自动化测试用例；将数据库密码、邮箱授权码、AI 密钥等"
     "敏感配置迁移至环境变量以提升部署安全性。总体而言，本次实习让小组完整经历了需求分析、系统设计、编码实现、"
     "联调测试到文档整理的软件工程全流程，在技术能力、工程规范与团队协作方面均获得了明显提升。")

# ============================================================
# 附录
# ============================================================
h1("附录")

h2("附录 A  用户端主要接口清单")
add_table(
    ["模块", "接口", "说明"],
    [
        ["用户", "POST /users/register", "注册"],
        ["用户", "POST /users/authenticate", "登录"],
        ["用户", "GET /users/{id}/profile", "用户公开资料"],
        ["资料", "PUT /api/me/profile", "更新个人资料"],
        ["分类", "GET /categories", "分类列表"],
        ["商品", "GET /products、GET /products/{id}", "列表/搜索、详情"],
        ["商品", "POST /products、PUT /products/{id}", "发布、编辑"],
        ["商品", "PUT /products/{id}/status、GET /products/my", "改状态、我的商品"],
        ["文件", "POST /files/upload", "图片上传"],
        ["订单", "POST /orders、GET /orders/{orderId}", "创建、详情"],
        ["订单", "GET /orders/my-purchases、/my-sales", "我的购买/出售"],
        ["订单", "PUT /orders/{orderId}/status", "更新订单状态"],
        ["评价", "POST /api/orders/{orderId}/ratings", "提交评价"],
        ["收藏", "POST/DELETE /api/products/{id}/favorite", "收藏/取消"],
        ["私信", "POST /api/messages、GET /api/messages/conversations", "发送、会话列表"],
        ["通知", "GET /api/notifications、/unread-count", "通知列表、未读数"],
        ["求购", "POST /demands、GET /demands/my", "发布、我的求购"],
        ["AI", "POST /ai/publish/suggest、/ai/price/suggest", "文案与价格建议"],
    ],
    widths=[2.2, 8.8, 5.0],
)

h2("附录 B  管理端主要接口清单")
add_table(
    ["模块", "接口", "说明"],
    [
        ["用户管理", "GET/POST/PUT/DELETE /admin/users、/status", "用户增删改查与状态"],
        ["商品管理", "GET/POST/PUT/DELETE /admin/products、/status", "商品增删改查与状态"],
        ["订单管理", "GET/POST/PUT/DELETE /admin/orders", "订单增删改查"],
        ["订单发货", "PUT /admin/orders/{id}/ship", "后台发货"],
        ["配送管理", "GET /admin/delivery/stats、/export", "配送统计、导出"],
        ["交易地点", "GET/POST/PUT/DELETE /admin/locations", "地点管理"],
        ["数据统计", "GET /admin/stats/summary 等", "汇总、注册、增长、订单/商品趋势"],
    ],
    widths=[2.6, 8.4, 5.0],
)

h2("附录 C  核心数据表清单")
add_table(
    ["数据表", "说明", "数据表", "说明"],
    [
        ["user", "用户", "notifications", "通知"],
        ["category", "分类", "meetup_location", "交易地点"],
        ["product", "商品", "user_addresses", "收货地址"],
        ["product_images", "商品图片", "product_demands", "求购需求"],
        ["orders", "订单", "product_embeddings", "商品向量"],
        ["ratings", "评价", "product_risks", "风险记录"],
        ["messages", "私信", "favorites", "收藏"],
    ],
    widths=[3.6, 4.4, 3.6, 4.4],
)

h2("附录 D  运行与启动说明")
bullets([
    "数据库初始化：创建 campus_trade 数据库并导入 campus-trade-api/src/main/resources/campus_trade.sql；",
    "后端启动：配置 application.yml（数据库、Redis、邮箱、AI Key）后，在 campus-trade-api 执行 mvn spring-boot:run；",
    "用户端启动：在 campus-trade-web 执行 npm install 与 npm run dev（http://localhost:5173）；",
    "管理端启动：在 campus-trade-admin 执行 npm install 与 npm run dev（http://localhost:8000）；",
    "默认管理员：admin / admin123，首次启动自动创建。",
])

para()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("—— 报告完 ——"); set_cn_font(r, FONT_CN, 12)

doc.save(OUT)
print("SAVED:", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
